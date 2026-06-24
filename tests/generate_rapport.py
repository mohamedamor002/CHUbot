"""
Génère rapport_stage.docx — rapport de stage CHUbot complet en Word.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Couleurs CHU ──────────────────────────────────────────────────────────────
BLUE_CHU    = RGBColor(0x1B, 0x4F, 0x9C)   # bleu CHU
BLUE_LIGHT  = RGBColor(0x2E, 0x86, 0xC1)   # bleu secondaire
GREY_DARK   = RGBColor(0x2C, 0x3E, 0x50)   # gris foncé texte
GREY_LIGHT  = RGBColor(0xF2, 0xF3, 0xF4)   # fond tableaux
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_OK    = RGBColor(0x1E, 0x8B, 0x4C)
ORANGE_WARN = RGBColor(0xE6, 0x7E, 0x22)
RED_ERR     = RGBColor(0xC0, 0x39, 0x2B)


# ── Helpers ───────────────────────────────────────────────────────────────────

def rgb_hex(color: RGBColor) -> str:
    return str(color)  # RGBColor.__str__ returns the hex string e.g. "1B4F9C"

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = rgb_hex(color)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CHUbot — Rapport de stage — CHU d'Angers — DRH    |    Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    add_page_number(p)

def h1(doc, text, page_break_before=True):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLUE_CHU
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1B4F9C")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BLUE_LIGHT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p

def h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = GREY_DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def bullet(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1 + level * 0.5)
        p.paragraph_format.space_after = Pt(3)

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F6F7")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x17, 0x20, 0x2A)
    return p

def info_box(doc, text, color_hex="E8F4FD", border_hex="2E86C1"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), border_hex)
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    return p

def make_table(doc, headers, rows, header_bg=BLUE_CHU, header_fg=WHITE,
               alt_rows=True, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = header_fg
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], header_bg)

    # Données
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = RGBColor(0xF2, 0xF3, 0xF4) if (alt_rows and r_idx % 2 == 0) else WHITE
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            if isinstance(val, tuple):
                text, bold, color = val
            else:
                text, bold, color = str(val), False, None
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            set_cell_bg(row_cells[c_idx], bg)

    # Largeurs colonnes
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def progress_bar(doc, label, score, target=None, note=""):
    """Affiche une barre de progression textuelle."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    filled = int(score * 20)
    bar = "█" * filled + "░" * (20 - filled)

    run1 = p.add_run(f"{label:<35}")
    run1.font.name = "Courier New"
    run1.font.size = Pt(9)
    run1.font.bold = True

    color = GREEN_OK if score >= 0.75 else (ORANGE_WARN if score >= 0.5 else RED_ERR)
    run2 = p.add_run(bar)
    run2.font.name = "Courier New"
    run2.font.size = Pt(9)
    run2.font.color.rgb = color

    run3 = p.add_run(f"  {score:.3f}/1.000")
    run3.font.name = "Courier New"
    run3.font.size = Pt(9)
    run3.font.bold = True
    run3.font.color.rgb = color

    if target:
        run4 = p.add_run(f"  (cible: {target})")
        run4.font.name = "Courier New"
        run4.font.size = Pt(8)
        run4.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    if note:
        p2 = doc.add_paragraph()
        r = p2.add_run(f"    → {note}")
        r.font.size = Pt(8.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.space_before = Pt(0)


# ── DOCUMENT ──────────────────────────────────────────────────────────────────

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Style Normal
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = GREY_DARK

add_footer(doc)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════

# Bandeau bleu supérieur (simulation via tableau 1×1)
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t.rows[0].cells[0]
set_cell_bg(cell, BLUE_CHU)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RAPPORT DE STAGE")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = WHITE
r.font.name = "Calibri"
cell.height = Cm(2.5)
doc.add_paragraph()

# Titre principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Développement d'un Chatbot RH Intelligent\npour le CHU d'Angers")
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = BLUE_CHU
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Basé sur une architecture RAG (Retrieval-Augmented Generation)")
r.font.size = Pt(13)
r.font.italic = True
r.font.color.rgb = BLUE_LIGHT
p.paragraph_format.space_after = Pt(30)

# Ligne de séparation
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), "1B4F9C")
pBdr.append(bottom); pPr.append(pBdr)

# Infos stage (tableau)
doc.add_paragraph()
t2 = doc.add_table(rows=5, cols=2)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
infos = [
    ("Établissement d'accueil", "CHU d'Angers"),
    ("Service", "Direction des Ressources Humaines (DRH)"),
    ("Période", "2025 – 2026"),
    ("Projet", "CHUbot v0.1.0"),
    ("Date du rapport", datetime.date.today().strftime("%d %B %Y")),
]
for i, (label, val) in enumerate(infos):
    cells = t2.rows[i].cells
    cells[0].text = ""
    r = cells[0].paragraphs[0].add_run(label)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLUE_CHU
    set_cell_bg(cells[0], RGBColor(0xEA, 0xF2, 0xFB))
    cells[1].text = ""
    r2 = cells[1].paragraphs[0].add_run(val)
    r2.font.size = Pt(10)
for row in t2.rows:
    row.cells[0].width = Cm(7)
    row.cells[1].width = Cm(9)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# RÉSUMÉ
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "1. Résumé", page_break_before=True)
info_box(doc,
    "Ce stage a consisté à concevoir et développer CHUbot, un assistant conversationnel "
    "intelligent destiné aux agents du CHU d'Angers. Le chatbot répond automatiquement "
    "aux questions relatives aux ressources humaines — congés, rémunération, formation, "
    "carrière, protection sociale — en s'appuyant sur la documentation officielle interne du CHU.",
    "E8F4FD", "2E86C1"
)
body(doc, "La solution repose sur une architecture RAG (Retrieval-Augmented Generation) "
     "entièrement locale : les documents RH sont vectorisés et stockés dans une base ChromaDB, "
     "puis un LLM open-source hébergé sur l'infrastructure CHU génère les réponses en citant "
     "ses sources.")
body(doc, "Un système de recherche hybride (BM25 + vectorielle), un reclassement sémantique "
     "(FlashRank) et des gardes anti-hallucination garantissent la fiabilité des réponses. "
     "L'application se présente sous forme d'une fenêtre flottante desktop (React + Tauri) "
     "toujours visible sur le poste de travail.")

h3(doc, "Mots-clés")
body(doc, "RAG • LLM • Ollama • LangChain • ChromaDB • FastAPI • Tauri • RH • chatbot • NLP • "
     "Python • React • PostgreSQL")


# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "2. Introduction")
body(doc, "La transformation numérique des services publics hospitaliers représente un enjeu "
     "majeur. Les agents d'un Centre Hospitalier Universitaire sont quotidiennement confrontés "
     "à des questions administratives complexes : calcul de congés, modalités de remboursement "
     "de transport, procédures d'avancement, droits liés au statut de fonctionnaire hospitalier…")
body(doc, "Ces questions, souvent récurrentes, mobilisent un temps précieux des gestionnaires RH. "
     "En parallèle, les agents se retrouvent parfois dans l'incapacité de trouver rapidement "
     "l'information dans la masse documentaire mise à leur disposition : notes d'information, "
     "guides, formulaires, sites institutionnels.")
body(doc, "L'émergence des modèles de langage de grande taille (LLM) ouvre une nouvelle voie : "
     "permettre à un agent de poser sa question en langage naturel et recevoir une réponse "
     "précise, sourcée et instantanée. C'est l'objectif de CHUbot.")
body(doc, "Ce rapport détaille l'ensemble du travail réalisé durant le stage : de l'analyse "
     "des besoins à la mise en production, en passant par les choix d'architecture, les "
     "développements techniques et l'évaluation des performances.")


# ══════════════════════════════════════════════════════════════════════════════
# ORGANISME D'ACCUEIL
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "3. Présentation de l'organisme d'accueil")
h2(doc, "3.1  Le CHU d'Angers")
body(doc, "Le Centre Hospitalier Universitaire d'Angers est l'un des établissements de santé "
     "les plus importants du Grand Ouest français. Il assure des missions de soins, "
     "d'enseignement médical et de recherche.")

make_table(doc,
    ["Indicateur", "Valeur"],
    [
        ("Nombre de lits", "~1 700"),
        ("Personnel médical et non-médical", "~7 000 agents"),
        ("Budget annuel", "> 600 M€"),
        ("Statut", "Établissement Public de Santé (EPCSM)"),
    ],
    col_widths=[8, 8]
)

h2(doc, "3.2  La Direction des Ressources Humaines")
body(doc, "La DRH du CHU d'Angers est organisée en 5 services couvrant l'ensemble du cycle "
     "de vie des agents :")
make_table(doc,
    ["Service", "Périmètre principal"],
    [
        ("Recrutement & Gestion des Effectifs",     "Concours, mobilité interne, mutations, détachements"),
        ("Service des Rémunérations",                "Paie, primes, bulletins, RIB, remboursements transport"),
        ("Gestion du Temps de Travail (GTT)",        "Congés, RTT, CET, absences, astreintes, missions"),
        ("Parcours Professionnels & Accompagnement", "Formation, CPF, ANFH, handicap (RQTH), entretien pro."),
        ("Service des Carrières",                    "Avancement, retraite, notation, positions statutaires"),
    ],
    col_widths=[7.5, 8.5]
)

h3(doc, "Volume documentaire géré")
bullet(doc, [
    "Notes d'information internes (NI) : publications régulières",
    "Formulaires administratifs : ~20 types",
    "Sites web institutionnels : 12 URLs référencées (CHU, ANFH, CGOS, CNRACL…)",
    "Guides thématiques : primes, temps de travail, retraite, formation…",
])


# ══════════════════════════════════════════════════════════════════════════════
# ANALYSE DU BESOIN
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "4. Analyse du besoin et problématique")
h2(doc, "4.1  Problèmes identifiés")
body(doc, "Une analyse des flux entrants vers la DRH a mis en évidence plusieurs pain points :")

make_table(doc,
    ["Problème", "Impact", "Cible"],
    [
        ("Questions RH répétitives (congés, RIB, primes)", "Temps mobilisé gestionnaires", "Agents"),
        ("Documentation dispersée et difficile à trouver", "Perte de temps agents", "Agents"),
        ("Horaires DRH limitées",                          "Pas de réponse hors heures ouverture", "Agents"),
        ("Mise à jour fréquente des documents",            "Risque d'information obsolète", "DRH"),
    ],
    col_widths=[6.5, 5, 4.5]
)

h2(doc, "4.2  Répartition des questions par thème")
body(doc, "Après analyse de la documentation et des retours métier, 105 questions types ont été "
     "identifiées et structurées en 12 catégories :", bold=True)
doc.add_paragraph()

# Graphique en barres (texte Courier)
p_title = doc.add_paragraph()
r = p_title.add_run("RÉPARTITION DES 105 QUESTIONS PAR CATÉGORIE")
r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = BLUE_CHU
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(6)

categories = [
    ("Rémunération",                         31, 30),
    ("Carrière",                             16, 15),
    ("Congés et absences",                   15, 14),
    ("Formation",                            12, 11),
    ("Contacts utiles",                      11, 10),
    ("Protection sociale",                    6,  6),
    ("Entretien professionnel",               4,  4),
    ("Temps de travail",                      4,  4),
    ("Accidents et maladies prof.",           2,  2),
    ("CGOS et action sociale",                2,  2),
    ("Crèches",                               1,  1),
    ("Retraite",                              1,  1),
]
for cat, n, pct in categories:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    bar_len = int(n / 31 * 30)
    bar = "█" * bar_len
    run1 = p.add_run(f"  {cat:<40}")
    run1.font.name = "Courier New"; run1.font.size = Pt(9)
    run2 = p.add_run(bar)
    run2.font.name = "Courier New"; run2.font.size = Pt(9)
    run2.font.color.rgb = BLUE_CHU
    run3 = p.add_run(f"  {n:>3} ({pct}%)")
    run3.font.name = "Courier New"; run3.font.size = Pt(9)
    run3.font.bold = True
doc.add_paragraph()

h2(doc, "4.3  Contraintes spécifiques au contexte hospitalier")
make_table(doc,
    ["Contrainte", "Exigence", "Solution retenue"],
    [
        ("Confidentialité données", "Aucune donnée sur le cloud",        "LLM et embeddings 100 % locaux"),
        ("Infrastructure existante", "Serveur Ollama interne CHU",       "Connexion https://llm.chu-angers.fr"),
        ("RGPD",                    "Pas de données personnelles",        "Filtre sujets sensibles + escalade"),
        ("Fiabilité",               "Aucune réponse inventée",           "Garde anti-hallucination + citations"),
        ("Déploiement",             "Poste Windows, pas de droits admin", "Tauri + PyInstaller (app autonome)"),
    ],
    col_widths=[4, 5, 7]
)

h2(doc, "4.4  Objectifs du projet")
bullet(doc, [
    "Répondre automatiquement aux questions RH fréquentes en langage naturel",
    "Citer les sources documentaires pour chaque réponse",
    "Ne jamais inventer d'information réglementaire (anti-hallucination)",
    "Détecter les cas complexes nécessitant l'intervention d'un gestionnaire",
    "S'intégrer dans le quotidien des agents (fenêtre toujours visible)",
    "Permettre le suivi de la qualité et des usages via des métriques",
])


# ══════════════════════════════════════════════════════════════════════════════
# ÉTAT DE L'ART
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "5. État de l'art — Les chatbots RAG")
h2(doc, "5.1  Pourquoi RAG et non un fine-tuning ?")
body(doc, "Deux approches majeures existent pour spécialiser un LLM sur un domaine métier. "
     "Le tableau ci-dessous compare les deux approches :")
make_table(doc,
    ["Critère", "Fine-tuning", "RAG (approche retenue)"],
    [
        ("Coût de mise en œuvre",      "Élevé (GPU, données d'entraînement)", ("Faible — documents existants suffisent", True, GREEN_OK)),
        ("Risque d'hallucination",     "Élevé",                               ("Faible — réponses ancrées dans les docs", True, GREEN_OK)),
        ("Mise à jour",                "Réentraînement complet",              ("Réindexation en quelques minutes", True, GREEN_OK)),
        ("Connaissance figée",         "Oui — gelée à l'entraînement",       ("Non — base documentaire évolutive", True, GREEN_OK)),
        ("Vitesse d'inférence",        ("Meilleure", True, GREEN_OK),         "Légère latence de récupération"),
        ("Adaptation au contexte CHU", "Difficile sans données annotées",     ("Directe — les docs RH sont la source", True, GREEN_OK)),
    ],
    col_widths=[5, 5.5, 5.5]
)
info_box(doc, "Conclusion : le RAG est clairement l'approche adaptée pour un contexte "
    "documentaire évolutif comme la DRH d'un CHU.", "E9F7EF", "1E8B4C")

h2(doc, "5.2  Fonctionnement du RAG")
body(doc, "Le RAG (Retrieval-Augmented Generation) opère en deux phases distinctes :")
bullet(doc, [
    "Phase 1 — Récupération (Retrieval) : recherche des documents les plus pertinents "
    "dans la base vectorielle à partir de la question de l'utilisateur.",
    "Phase 2 — Génération (Generation) : le LLM génère une réponse en se basant "
    "UNIQUEMENT sur les documents récupérés, ce qui élimine les hallucinations.",
])

h2(doc, "5.3  Technologies évaluées et retenues")
make_table(doc,
    ["Composant", "Alternatives évaluées", "Choix retenu", "Justification"],
    [
        ("LLM",          "GPT-4, Claude, Mistral",    "llama3.1 via Ollama",       "Local, gratuit, performant en français"),
        ("Embeddings",   "OpenAI, Cohere, HuggingFace","nomic-embed-text",          "Local, 768 dims, multilingue"),
        ("Vector DB",    "FAISS, Weaviate, Pinecone",  "ChromaDB",                  "Persistant, local, Python natif"),
        ("Framework",    "Haystack, LlamaIndex",       "LangChain",                 "Maturité, flexibilité, communauté"),
        ("Reranker",     "CohereRerank, Colbert",      "FlashRank",                 "CPU-only, aucune clé API requise"),
        ("Frontend",     "Electron, webapp",           "Tauri + React",             "Léger (Rust), natif Windows"),
        ("Base de données","SQLite, MongoDB",          "PostgreSQL",                "Robustesse, async, requêtes analytiques"),
    ],
    col_widths=[3, 4.5, 4, 4.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# ARCHITECTURE GÉNÉRALE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "6. Architecture générale du système")
h2(doc, "6.1  Vue d'ensemble")
body(doc, "Le système CHUbot se décompose en trois grandes couches :")
make_table(doc,
    ["Couche", "Technologies", "Rôle"],
    [
        ("Présentation",  "Tauri 2 + React 18 + Tailwind",  "Interface flottante desktop Windows"),
        ("Service",       "FastAPI + uvicorn (port 8765)",   "API REST async, streaming, analytics"),
        ("Données/IA",    "ChromaDB + PostgreSQL + Ollama",  "Vecteurs, conversations, LLM"),
    ],
    col_widths=[4, 6, 6]
)

body(doc, "Architecture de déploiement :")
code_block(doc,
"Poste Agent (Windows)\n"
"├── Interface Tauri  (React + Tailwind, fenêtre 80×80 ou 380×620)\n"
"│       ↕ HTTP localhost:8765\n"
"├── Backend FastAPI  (uvicorn, Python 3.11+)\n"
"│   ├── PostgreSQL  (conversations, métriques, feedback)\n"
"│   └── ChromaDB    (embeddings vectoriels, ./data/indexes)\n"
"│\n"
"└── Réseau interne CHU  →  https://llm.chu-angers.fr/ollama\n"
"                              ├── llama3.1:latest  (génération)\n"
"                              └── nomic-embed-text (vectorisation)"
)

h2(doc, "6.2  Flux de traitement d'une question")
body(doc, "Voici le déroulement complet d'une question posée par un agent :")
make_table(doc,
    ["Étape", "Composant", "Action"],
    [
        ("1", "preprocessor.py",  "Expansion des acronymes RH (RTT, CET, FMD…)"),
        ("2", "query_rewriter.py","Réécriture de la question par LLM (mots-clés optimaux)"),
        ("3", "departments.py",   "Détection du sous-département concerné"),
        ("4", "chain.py",         "Recherche BM25 (40%) + vectorielle (60%)"),
        ("5", "reranker.py",      "Reranking FlashRank → top-3 documents"),
        ("6", "chain.py",         "Garde : documents trouvés ? (sinon message fixe)"),
        ("7", "Ollama LLM",       "Génération de la réponse (streaming token/token)"),
        ("8", "database.py",      "Sauvegarde conversation + métriques en base"),
    ],
    col_widths=[1.5, 4.5, 10]
)

h2(doc, "6.3  Stack technologique complète")
make_table(doc,
    ["Couche", "Technologie", "Version", "Détails"],
    [
        ("Présentation", "Tauri",          "2.x",    "Runtime desktop natif Rust, installateur NSIS/MSI"),
        ("Présentation", "React",          "18.3.1", "Composants UI, gestion état"),
        ("Présentation", "Tailwind CSS",   "3.4.6",  "Design system utilitaire"),
        ("Présentation", "Vite",           "5.3.4",  "Bundler et dev server"),
        ("Présentation", "Axios",          "1.7.2",  "Client HTTP + streaming"),
        ("API",          "FastAPI",        "0.11+",  "Framework REST async"),
        ("API",          "uvicorn",        "latest", "Serveur ASGI"),
        ("API",          "Pydantic v2",    "2.x",    "Validation des schémas"),
        ("Métier (RAG)", "LangChain",      "0.3+",   "Orchestration LLM/RAG"),
        ("Métier (RAG)", "Ollama",         "latest", "LLM + embeddings locaux"),
        ("Métier (RAG)", "ChromaDB",       "0.6+",   "Base vectorielle persistante"),
        ("Métier (RAG)", "FlashRank",      "latest", "Reranker cross-encoder CPU"),
        ("Ingestion",    "PyMuPDF",        "latest", "Extraction PDF"),
        ("Ingestion",    "python-docx",    "1.2.0",  "Extraction Word"),
        ("Ingestion",    "openpyxl",       "latest", "Extraction Excel"),
        ("Ingestion",    "trafilatura",    "latest", "Scraping HTML statique"),
        ("Ingestion",    "Playwright",     "latest", "Scraping JavaScript"),
        ("Persistance",  "PostgreSQL",     "15+",    "SGBD relationnel"),
        ("Persistance",  "SQLAlchemy",     "2.x",    "ORM async"),
        ("Persistance",  "asyncpg",        "latest", "Driver PostgreSQL async"),
    ],
    col_widths=[3, 3.5, 2, 7.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE D'INGESTION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "7. Pipeline d'ingestion documentaire")
h2(doc, "7.1  Vue d'ensemble")
body(doc, "L'ingestion est la phase qui transforme des documents bruts en vecteurs interrogeables. "
     "Le pipeline se déroule en 5 étapes séquentielles :")

make_table(doc,
    ["Étape", "Module", "Entrée → Sortie"],
    [
        ("1 — Chargement",       "loaders.py",  "PDF / DOCX / Excel / HTML  →  texte brut"),
        ("2 — Nettoyage",        "cleaner.py",  "Texte brut  →  texte normalisé"),
        ("3 — Découpage",        "splitter.py", "Texte long  →  chunks 2000 chars avec en-tête contextuel"),
        ("4 — Déduplication",    "indexer.py",  "Chunks  →  liste unique (MD5)"),
        ("5 — Vectorisation",    "indexer.py",  "Chunks  →  embeddings ChromaDB (nomic-embed-text)"),
    ],
    col_widths=[4, 3.5, 8.5]
)

h2(doc, "7.2  Formats supportés")
make_table(doc,
    ["Format", "Librairie", "Particularités"],
    [
        ("PDF",          "PyMuPDF (fitz)",   "Extraction page par page, gestion colonnes, métadonnées"),
        ("DOCX / DOC",   "python-docx",      "Paragraphes + tableaux convertis en texte structuré"),
        ("XLSX / XLS",   "openpyxl",         "Chaque ligne : clé: valeur avec en-tête de colonne"),
        ("HTML statique","trafilatura",       "Extraction contenu principal, suppression nav/pied"),
        ("HTML dynamique","Playwright",       "Rendu JavaScript complet avant extraction"),
        ("Site web",     "BeautifulSoup",     "Crawl jusqu'à 30 pages, liens relatifs résolus"),
    ],
    col_widths=[3.5, 4, 8.5]
)

h2(doc, "7.3  En-tête contextuel des chunks")
body(doc, "Chaque chunk est précédé d'un en-tête contextuel qui améliore significativement "
     "la précision des embeddings :")
code_block(doc,
"[Source: 2025-46 NI Forfait mobilités durables | Section: Conditions d'éligibilité]\n"
"Les agents ayant recours au covoiturage, au vélo, à la trottinette électrique\n"
"ou à d'autres modes de transport dit « doux » pour leurs déplacements entre\n"
"leur résidence habituelle et leur lieu de travail peuvent bénéficier du\n"
"Forfait Mobilités Durables (FMD)...\n"
"(chunk 3/8 | sous-département: Rémunérations)"
)

h2(doc, "7.4  Paramètres de chunking")
make_table(doc,
    ["Paramètre", "Valeur", "Justification"],
    [
        ("Taille du chunk",   "2 000 caractères (~500 tokens)", "Optimal pour llama3.1 (fenêtre 4K tokens)"),
        ("Overlap",           "400 caractères (20%)",           "Évite la coupure des phrases frontières"),
        ("Séparateurs",       "\\n\\n, \\n, . , espace",         "Découpe aux frontières naturelles du texte"),
        ("Métadonnées",       "source, section, chunk_index",   "Traçabilité et filtrage par sous-département"),
    ],
    col_widths=[4, 5, 7]
)

h2(doc, "7.5  Interface CLI de l'indexeur")
code_block(doc,
"# Indexer un répertoire complet\n"
"python -m bot.ingestion.indexer --dir data/documents/\n\n"
"# Indexer un fichier unique\n"
"python -m bot.ingestion.indexer --file data/documents/guide_primes.pdf\n\n"
"# Indexer les 12 URLs documentaires (CHU, ANFH, CGOS, CNRACL...)\n"
"python -m bot.ingestion.indexer --urls\n\n"
"# Tout indexer avec reset préalable\n"
"python -m bot.ingestion.indexer --all --reset"
)


# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE RAG
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "8. Pipeline RAG — Retrieval-Augmented Generation")
h2(doc, "8.1  Prétraitement de la question")
body(doc, "Avant toute recherche, la question de l'agent subit deux transformations successives.")

h3(doc, "A — Expansion des acronymes RH")
body(doc, "18 acronymes spécifiques à la fonction publique hospitalière sont développés "
     "automatiquement par substitution regex :")
make_table(doc,
    ["Acronyme", "Développement"],
    [
        ("RTT",    "Réduction du Temps de Travail (RTT)"),
        ("CET",    "Compte Épargne-Temps (CET)"),
        ("CPF",    "Compte Personnel de Formation (CPF)"),
        ("FMD",    "Forfait Mobilités Durables (FMD)"),
        ("NBI",    "Nouvelle Bonification Indiciaire (NBI)"),
        ("RME",    "Reconnaissance de la Qualité de Travailleur Handicapé (RQTH/RME)"),
        ("ANFH",   "Association Nationale pour la Formation des Hospitaliers (ANFH)"),
        ("CGOS",   "Comité de Gestion des Œuvres Sociales (CGOS)"),
        ("CNRACL", "Caisse Nationale de Retraite des Agents des Collectivités Locales"),
        ("GTT",    "Gestion du Temps de Travail (GTT)"),
    ],
    col_widths=[3, 13]
)

h3(doc, "B — Réécriture de la question par LLM")
body(doc, "Un appel LLM (température 0.0) reformule la question en mots-clés optimaux pour la "
     "recherche documentaire :")
code_block(doc,
'Input  : "Comment prendre mon CET ?"\n\n'
'Prompt : "Tu es un expert RH. Réécris cette question en 5-10 mots-clés précis\n'
'          pour une recherche documentaire. Évite les confusions sémantiques."\n\n'
'Output : "utilisation compte épargne-temps CET congés hospitalier déblocage\n'
'          procédure service RH"'
)

h2(doc, "8.2  Recherche hybride avec Reciprocal Rank Fusion")
body(doc, "La recherche hybride combine deux approches complémentaires fusionnées par "
     "Reciprocal Rank Fusion (RRF) :")
make_table(doc,
    ["Méthode", "Poids", "Avantage", "Limite"],
    [
        ("BM25 (lexicale)",      "40%",  "Excellent sur termes exacts, acronymes", "Sensible au vocabulaire exact"),
        ("Vectorielle (sémant.)", "60%", "Comprend le sens, synonymes",            "Peut manquer les termes précis"),
        ("RRF (fusion)",          "—",   "Combine les deux, robuste",              "—"),
    ],
    col_widths=[4, 2, 6, 4]
)
body(doc, "Formule RRF : score(d) = Σᵢ 1/(k + rankᵢ(d))  avec k=60 (paramètre de stabilité)")

h2(doc, "8.3  Reranking sémantique FlashRank")
body(doc, "Après la fusion RRF, les documents sont reclassés par un cross-encoder "
     "(ms-marco-MiniLM-L-12-v2) qui évalue la pertinence de chaque document vis-à-vis "
     "de la question :")
make_table(doc,
    ["Paramètre", "Valeur"],
    [
        ("Modèle",          "ms-marco-MiniLM-L-12-v2"),
        ("Taille",          "~130 MB (CPU-only, pas de GPU requis)"),
        ("Seuil de score",  "0.05 (documents en-dessous exclus)"),
        ("Top-K conservés", "3 documents (après reranking)"),
    ],
    col_widths=[5, 11]
)

h2(doc, "8.4  Filtre par sous-département")
body(doc, "La détection du sous-département permet de filtrer ChromaDB sur les métadonnées "
     "avant la recherche, améliorant la précision :")
make_table(doc,
    ["Département détecté", "Mots-clés déclencheurs (exemples)"],
    [
        ("Rémunérations",           "primes, paie, bulletin, RIB, cotisation, salaire"),
        ("GTT",                     "congés, RTT, CET, absence, astreinte, mission"),
        ("Recrutement",             "concours, mobilité, mutation, détachement"),
        ("Carrières",               "avancement, retraite, notation, statut, échelon"),
        ("Parcours Professionnels", "formation, CPF, ANFH, handicap, entretien professionnel"),
    ],
    col_widths=[5, 11]
)
body(doc, "Si le filtre retourne 0 résultats → fallback automatique vers une recherche globale "
     "sans filtre.", italic=True)

h2(doc, "8.5  Garde anti-hallucination")
body(doc, "Si FlashRank ne retourne aucun document dépassant le seuil, le LLM n'est "
     "JAMAIS appelé. Une réponse fixe et honnête est retournée :", bold=True)
code_block(doc,
'reranked = [d for d in docs if d.metadata.get("_rerank_score", 0) >= 0.05]\n\n'
'if not reranked:\n'
'    return (\n'
'        "Je n\'ai pas trouvé d\'information dans la documentation disponible\\n"\n'
'        "pour répondre à votre question. Je vous invite à contacter\\n"\n'
'        "directement le service RH concerné."\n'
'    )\n\n'
'# Le LLM ne répond QUE sur la base du contexte fourni'
)

h2(doc, "8.6  Prompt système et règles de génération")
info_box(doc,
    "Règles absolues du prompt système :\n"
    "1. Réponses basées UNIQUEMENT sur les documents fournis\n"
    "2. Jamais de données individuelles (salaire, contrat, données personnelles)\n"
    "3. Toujours citer la source documentaire\n"
    "4. Si information absente : le dire explicitement et orienter vers le service compétent\n"
    "5. Répondre en français, clair et concis\n"
    "6. Mentionner les contacts d'escalade si nécessaire",
    "FEF9E7", "E67E22"
)

h2(doc, "8.7  Contacts d'escalade automatiques")
make_table(doc,
    ["Trigger détecté", "Service", "Email"],
    [
        ("Accident de travail",     "Service RH — AT/MP",      "rh-at@chu-angers.fr"),
        ("Situation carrière complex.", "Service Carrières",    "rh-carrieres@chu-angers.fr"),
        ("Litige rémunération",     "Service Rémunérations",   "rh-remuneration@chu-angers.fr"),
        ("Formation ANFH",          "Service Formation",       "rh-formation@chu-angers.fr"),
    ],
    col_widths=[5, 5, 6]
)


# ══════════════════════════════════════════════════════════════════════════════
# BASE DE DONNÉES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "9. Base de données et persistance")
h2(doc, "9.1  Schéma relationnel — 6 tables")
body(doc, "La base PostgreSQL assure la persistance complète des échanges et des métriques :")

make_table(doc,
    ["Table", "Colonnes principales", "Rôle"],
    [
        ("users",               "id, username, email, hashed_password, role",         "Comptes agents"),
        ("conversation_sessions","id, user_id, title, created_at, updated_at",        "Sessions de chat"),
        ("messages",            "id, session_id, role, content, created_at",          "Messages Q/R"),
        ("feedback",            "id, message_id, is_helpful, rating (1-5)",           "Évaluations agents"),
        ("message_metrics",     "id, message_id, response_time_ms, is_covered, has_escalation", "KPIs techniques"),
        ("indexed_documents",   "id, filename, file_type, chunks_indexed, indexed_at","Traçabilité ingestion"),
    ],
    col_widths=[4, 7, 5]
)

h2(doc, "9.2  Architecture asynchrone SQLAlchemy")
body(doc, "Le driver asyncpg permet de ne pas bloquer l'event loop FastAPI pendant les "
     "requêtes SQL (importantes lors du streaming LLM) :")
code_block(doc,
"# Session par requête HTTP (Dependency Injection FastAPI)\n"
"async def get_db() -> AsyncGenerator[AsyncSession, None]:\n"
"    async with AsyncSessionLocal() as session:\n"
"        try:\n"
"            yield session\n"
"            await session.commit()\n"
"        except Exception:\n"
"            await session.rollback()\n"
"            raise"
)


# ══════════════════════════════════════════════════════════════════════════════
# API REST
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "10. API REST — Couche de service")
h2(doc, "10.1  Carte des 11 endpoints")
make_table(doc,
    ["Méthode", "Endpoint", "Description"],
    [
        ("POST", "/api/v1/chat",                  "Réponse complète (non-streaming)"),
        ("POST", "/api/v1/chat/stream",            "Réponse streaming token par token"),
        ("GET",  "/api/v1/chat/sessions",          "50 sessions les plus récentes"),
        ("GET",  "/api/v1/chat/sessions/{id}",     "Historique complet d'une session"),
        ("POST", "/api/v1/documents/upload",       "Upload + indexation immédiate d'un fichier"),
        ("GET",  "/api/v1/documents/",             "Lister les documents indexés"),
        ("GET",  "/api/v1/documents/file/{name}",  "Télécharger un document source"),
        ("POST", "/api/v1/feedback",               "Soumettre un avis (👍/👎 + note 1-5)"),
        ("GET",  "/api/v1/analytics/kpis",         "Tableau de bord des métriques agrégées"),
        ("GET",  "/health",                        "Statut du serveur"),
    ],
    col_widths=[2, 6, 8]
)

h2(doc, "10.2  Headers du streaming")
body(doc, "L'endpoint /chat/stream retourne des headers HTTP personnalisés pour permettre "
     "au frontend de lier la session et les sources :")
code_block(doc,
"X-Session-ID  : uuid-de-la-session     → Maintenir/créer la session\n"
"X-Message-ID  : uuid-du-message        → Pour feedback ultérieur\n"
"X-Sources     : [{filename, url, type}] → Documents sources à afficher\n"
"Content-Type  : text/plain              → Tokens text au fil de l'eau"
)

h2(doc, "10.3  KPIs analytiques")
make_table(doc,
    ["KPI", "Calcul", "Cible"],
    [
        ("Taux réponses utiles",  "% feedback is_helpful = true",        "≥ 80%"),
        ("Taux couverture docs",  "% messages is_covered = true",        "≥ 85%"),
        ("Taux d'escalade",       "% messages has_escalation = true",    "≤ 20%"),
        ("Score satisfaction",    "Moyenne feedback.rating (1-5)",       "≥ 4.0/5"),
        ("Temps réponse moyen",   "Moyenne response_time_ms",            "≤ 5 000 ms"),
    ],
    col_widths=[5, 6, 5]
)


# ══════════════════════════════════════════════════════════════════════════════
# INTERFACE UTILISATEUR
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "11. Interface utilisateur — Application Desktop")
h2(doc, "11.1  Concept UX : le chatbot flottant")
body(doc, "L'interface a été conçue pour s'intégrer au quotidien sans perturber le flux de "
     "travail de l'agent. Elle opère en deux états :")

make_table(doc,
    ["État", "Dimensions", "Comportement"],
    [
        ("Avatar", "80 × 80 px", "Icône médicale bleue flottante, animation pulse, draggable, toujours au premier plan"),
        ("Chat",   "380 × 620 px","Fenêtre de chat complète avec historique, streaming, sources, feedback"),
    ],
    col_widths=[2.5, 3.5, 10]
)
body(doc, "Transition : clic sur l'avatar → redimensionnement programmatique Tauri → ouverture du chat. "
     "Bouton réduire → retour à l'avatar.", italic=True)

h2(doc, "11.2  Arbre des composants React")
make_table(doc,
    ["Composant", "Rôle", "Détails techniques"],
    [
        ("App.jsx",           "Racine, gestion état global",    "Transitions Avatar ↔ Chat + resize Tauri API"),
        ("Avatar.jsx",        "Icône flottante 80×80",          "Croix médicale, animation ping, drag region"),
        ("FloatingChat.jsx",  "Conteneur chat 380×620",         "Nouvelle conversation, badge version"),
        ("WindowControls.jsx","Contrôles fenêtre",              "Min / Max / Close / Collapse (Windows-style)"),
        ("ChatWindow.jsx",    "Chat principal",                  "Message list, auto-scroll, questions suggérées"),
        ("ChatInput.jsx",     "Zone de saisie",                 "Auto-resize, Enter=Envoyer, max 2000 chars"),
        ("MessageBubble.jsx", "Bulle de message",                "ReactMarkdown, sources, boutons feedback"),
        ("api/client.js",     "Client HTTP",                    "Axios + streaming ReadableStream, timeout 120s"),
    ],
    col_widths=[4, 4, 8]
)

h2(doc, "11.3  Fenêtre frameless Tauri")
body(doc, "La fenêtre Tauri est configurée sans décoration Windows, transparente et toujours "
     "au premier plan. Le redimensionnement est entièrement géré par le code :")
code_block(doc,
"// Transition Avatar → Chat\n"
"await appWindow.setSize(new PhysicalSize(380, 620));\n"
"await appWindow.center();\n"
"setIsOpen(true);\n\n"
"// Transition Chat → Avatar\n"
"await appWindow.setSize(new PhysicalSize(80, 80));\n"
"setIsOpen(false);"
)

h2(doc, "11.4  Rendu Markdown et liens documents")
body(doc, "Les réponses du LLM sont rendues en Markdown enrichi par ReactMarkdown. "
     "Les liens vers des documents PDF/DOCX ouvrent le fichier via l'API Tauri Opener, "
     "les liens web s'ouvrent dans le navigateur système.")

h2(doc, "11.5  Gestion du feedback")
body(doc, "Après chaque réponse, l'agent peut évaluer en un clic :")
bullet(doc, [
    "👍 Réponse utile → POST /feedback {is_helpful: true}",
    "👎 Réponse non utile → POST /feedback {is_helpful: false}",
    "Les boutons sont désactivés après vote (pas de double vote)",
    "Les données alimentent les KPIs du tableau de bord analytique",
])


# ══════════════════════════════════════════════════════════════════════════════
# ÉVALUATION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "12. Évaluation et métriques de performance")
h2(doc, "12.1  Dataset d'évaluation — 105 questions")
body(doc, "Un jeu de données de 105 paires question/réponse attendue a été construit "
     "manuellement en collaboration avec les équipes RH :")
make_table(doc,
    ["Catégorie", "Nb questions", "Exemples de questions"],
    [
        ("Rémunération",                 "31", "Bulletins de salaire, Digiposte, primes, RIB"),
        ("Carrière",                     "16", "Avancement, notation, positions statutaires"),
        ("Congés et absences",           "15", "Congés annuels, RTT, CET, congé parental"),
        ("Formation",                    "12", "CPF, ANFH, études promotionnelles"),
        ("Contacts utiles",              "11", "Contacts DRH, services, emails"),
        ("Protection sociale",            "6", "Mutuelle, accidents de travail"),
        ("Entretien professionnel",        "4", "Préparation, suites, procédure"),
        ("Temps de travail",               "4", "Astreintes, missions, télétravail"),
        ("Accidents et maladies prof.",    "2", "Déclaration, droits"),
        ("CGOS et action sociale",         "2", "Avantages CGOS"),
        ("Crèches / Retraite",             "2", "Places crèche, CNRACL"),
    ],
    col_widths=[5.5, 3, 7.5]
)

h2(doc, "12.2  Métriques d'évaluation automatique")
body(doc, "L'évaluation utilise le même LLM local (llama3.1) comme juge, sans dépendance "
     "à des services externes payants (pas d'OpenAI) :")
make_table(doc,
    ["Métrique", "Question posée au LLM-juge", "Interprétation"],
    [
        ("Faithfulness\n(0 → 1)",        "La réponse est-elle basée UNIQUEMENT sur le contexte ?", "Mesure l'anti-hallucination"),
        ("Answer Correctness\n(0 → 1)", "La réponse est-elle correcte vs la vérité terrain ?",    "Mesure la justesse factuelle"),
        ("Context Relevance\n(0 → 1)", "Le contexte récupéré est-il pertinent ?",                  "Mesure la qualité du retrieval"),
    ],
    col_widths=[4, 7, 5]
)

h2(doc, "12.3  Résultats d'évaluation")
body(doc, "Évaluation réalisée sur 10 questions de la catégorie Rémunération :", bold=True)
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("SCORES GLOBAUX — ÉVALUATION CHUBOT  (10 questions, catégorie Rémunération)")
r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLUE_CHU
p.paragraph_format.space_after = Pt(8)

progress_bar(doc, "Faithfulness (anti-hallucination)",  0.810, "> 0.75",
             "Le chatbot base ses réponses sur les documents ✓")
progress_bar(doc, "Context Relevance (qualité retrieval)", 0.740, "> 0.70",
             "Les documents récupérés sont pertinents ✓")
progress_bar(doc, "Answer Correctness (justesse fact.)", 0.580, "> 0.70",
             "Marge de progression — cas sans doc source exact")
doc.add_paragraph()

h3(doc, "Distribution des scores de correctness")
make_table(doc,
    ["Tranche de score", "Nombre de questions", "Pourcentage", "Interprétation"],
    [
        ("≥ 0.8  (excellent)",     ("6", True, GREEN_OK),  ("60%", True, GREEN_OK),  "Réponse fidèle à la vérité terrain"),
        ("0.5 – 0.79 (partielle)", ("0", False, None),     ("0%",  False, None),     "—"),
        ("< 0.5  (insuffisante)",  ("4", True, ORANGE_WARN),("40%", True, ORANGE_WARN),"Doc source non indexé ou cas ambigu"),
    ],
    col_widths=[4, 3.5, 3, 5.5]
)

h2(doc, "12.4  Analyse qualitative — exemples")
make_table(doc,
    ["Question (ID)", "Faith.", "Correct.", "Analyse"],
    [
        ("QR001 — Accès coffre-fort numérique", "0.90", "0.80", "✓ Correcte, source citée. Manque URL adhésion Digiposte"),
        ("QR003 — Départ CHU et bulletins",     "0.90", "0.30", "✗ Répond 'scanner les bulletins' → incorrect. Digiposte est à vie"),
        ("QR005 — Confidentialité coffre-fort", "0.90", "0.30", "✗ Document source insuffisant sur ce sujet précis"),
        ("QR008 — Forfait mobilités durables",  "0.20", "0.30", "✗ NI pas encore indexée au moment du test → réponse générique"),
    ],
    col_widths=[5.5, 1.5, 1.5, 7.5]
)
info_box(doc,
    "Conclusion : les 4 cas avec score < 0.5 sont liés soit à des documents manquants dans l'index "
    "(NI pas encore indexée), soit à des informations implicites non documentées explicitement. "
    "Cela indique que la qualité du chatbot est directement liée à l'exhaustivité de la base documentaire.",
    "FEF9E7", "E67E22"
)

h2(doc, "12.5  Module RAGAS")
body(doc, "En complément, un module RAGAS (Automated Evaluation Framework) a été intégré "
     "pour des métriques standardisées dans la communauté RAG :")
make_table(doc,
    ["Métrique RAGAS", "Colonnes requises", "Description"],
    [
        ("Faithfulness",       "response, user_input, retrieved_contexts", "Fidélité au contexte"),
        ("Answer Correctness", "response, user_input, reference",           "Justesse factuelle"),
        ("Context Precision",  "user_input, reference, retrieved_contexts","Précision du retrieval"),
    ],
    col_widths=[4, 6.5, 5.5]
)
code_block(doc,
"# Lancement de l'évaluation avec RAGAS\n"
"python tests/evaluate.py --limit 10 --ragas\n\n"
"# Résultats exportés en CSV et JSON\n"
"# data/eval/results-YYYYMMDD-HHMMSS.csv"
)


# ══════════════════════════════════════════════════════════════════════════════
# DÉPLOIEMENT
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "13. Déploiement et packaging")
h2(doc, "13.1  Architecture de packaging")
body(doc, "L'objectif est de livrer un fichier d'installation Windows autonome (.exe ou .msi) "
     "sans prérequis sur le poste de l'agent :")
code_block(doc,
"Setup.exe / Setup.msi\n"
"├── Frontend Tauri\n"
"│   ├── Interface React compilée (dist/)\n"
"│   └── Runtime Rust + WebView2 embarqué\n"
"│\n"
"└── Backend PyInstaller bundle\n"
"    ├── Python 3.11 embarqué\n"
"    ├── Toutes les dépendances Python\n"
"    ├── chubot-backend.exe  (démarre sur port 8765)\n"
"    └── FlashRank model cache"
)

h2(doc, "13.2  Processus de build")
make_table(doc,
    ["Étape", "Commande", "Résultat"],
    [
        ("1 — Backend",  "pyinstaller chubot-backend.spec",  "binaries/chubot-backend.exe (~150 MB)"),
        ("2 — Frontend", "cd bot/ui && npm run tauri build", "CHUbot_0.1.0_x64-setup.exe + .msi"),
    ],
    col_widths=[2.5, 7, 6.5]
)

h2(doc, "13.3  Configuration Tauri")
make_table(doc,
    ["Paramètre tauri.conf.json", "Valeur", "Justification"],
    [
        ("decorations",  "false",  "Pas de bordure Windows → design flottant"),
        ("transparent",  "true",   "Fond transparent pour l'avatar"),
        ("alwaysOnTop",  "true",   "Toujours visible au-dessus des applications"),
        ("width/height", "80/80",  "Taille initiale avatar"),
        ("targets",      "nsis + msi", "Deux formats d'installateur Windows"),
    ],
    col_widths=[5, 3, 8]
)


# ══════════════════════════════════════════════════════════════════════════════
# DIFFICULTÉS RENCONTRÉES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "14. Difficultés rencontrées et solutions")

make_table(doc,
    ["Problème", "Symptôme", "Solution apportée"],
    [
        (
            "Conflits asyncio / RAGAS\n(Python 3.14)",
            "Crash lors cleanup event loop —\nnest_asyncio incompatible",
            "Extraction RAGAS hors boucle asyncio :\nexécution synchrone après loop.close()"
        ),
        (
            "PermissionError results.csv",
            "[Errno 13] fichier ouvert\ndans un autre programme",
            "Fallback automatique vers fichier\nhorodaté (results-YYYYMMDD-HHMMSS.csv)"
        ),
        (
            "Qualité extractions PDF numérisés",
            "Texte incohérent, artefacts\nhyphenation, en-têtes répétés",
            "Module cleaner.py : regex de\nnettoyage spécifiques aux docs CHU"
        ),
        (
            "Hallucinations hors-corpus",
            "LLM invente des réponses\nplausibles mais incorrectes",
            "Garde anti-hallucination : seuil\nFlashRank 0.05 → LLM jamais appelé"
        ),
        (
            "Port 8765 déjà utilisé",
            "[WinError 10048] adresse\nsocket déjà utilisée",
            "Identification et kill du process\nchubot-backend avant redémarrage"
        ),
        (
            "Latence cold start (>30s)",
            "Première requête lente :\nchargement modèles et index",
            "Pre-warming dans lifespan FastAPI :\nget_vector_store() + get_ranker() au démarrage"
        ),
        (
            "Questions hors-périmètre RH",
            "Small talk, demandes\nnon-RH envoyées au LLM",
            "Détection par patterns regex\n→ réponse fixe sans appel LLM"
        ),
    ],
    col_widths=[4.5, 4.5, 7]
)


# ══════════════════════════════════════════════════════════════════════════════
# PERSPECTIVES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "15. Perspectives d'évolution")
h2(doc, "15.1  Court terme (< 3 mois)")
bullet(doc, [
    "Authentification agents via SSO CHU pour personnaliser les réponses",
    "Watcher automatique sur data/documents/ pour réindexation en temps réel",
    "Amélioration du dataset à 200+ questions (couverture toutes catégories)",
    "Mode hors-ligne : cache des Q/R fréquentes sans réseau",
])

h2(doc, "15.2  Moyen terme (3 – 12 mois)")
bullet(doc, [
    "Tableau de bord manager : interface dédiée DRH pour les KPIs",
    "Apprentissage par feedback : utiliser les 👎 pour améliorer automatiquement",
    "Citations précises : pointer vers la page exacte du document source",
    "Intégration Sharepoint : indexation automatique de la GED CHU",
    "Support multi-langue (anglais, arabe) pour agents étrangers",
])

h2(doc, "15.3  Long terme")
bullet(doc, [
    "Agents spécialisés : un sous-LLM par département RH avec expertise ciblée",
    "Intégration SI-RH (Gestime) : réponses personnalisées (congés restants, solde CET...)",
    "Assistance multi-modale : analyser des formulaires photographiés",
    "Mesure de l'impact : réduction quantifiée des sollicitations DRH",
])


# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "16. Conclusion")
body(doc, "Ce stage au CHU d'Angers m'a permis de mener de bout en bout le développement "
     "d'un système RAG de niveau production, en répondant à de véritables contraintes "
     "métier et techniques.", size=11)

h2(doc, "16.1  Bilan des composants développés")
make_table(doc,
    ["Composant", "Statut"],
    [
        ("Pipeline d'ingestion (PDF, DOCX, Excel, Web)",           ("✅ Opérationnel", True, GREEN_OK)),
        ("Recherche hybride BM25 + vectorielle + RRF",             ("✅ Opérationnel", True, GREEN_OK)),
        ("Reranking FlashRank (cross-encoder CPU)",                 ("✅ Opérationnel", True, GREEN_OK)),
        ("Garde anti-hallucination systématique",                   ("✅ Opérationnel", True, GREEN_OK)),
        ("Streaming LLM token par token",                          ("✅ Opérationnel", True, GREEN_OK)),
        ("API REST FastAPI (11 endpoints)",                        ("✅ Opérationnel", True, GREEN_OK)),
        ("Persistance PostgreSQL async (6 tables)",                ("✅ Opérationnel", True, GREEN_OK)),
        ("Interface Tauri flottante React",                        ("✅ Opérationnel", True, GREEN_OK)),
        ("Système de feedback utilisateur",                        ("✅ Opérationnel", True, GREEN_OK)),
        ("Analytics KPI (5 métriques)",                           ("✅ Opérationnel", True, GREEN_OK)),
        ("Framework d'évaluation automatique (LLM-as-judge + RAGAS)", ("✅ Opérationnel", True, GREEN_OK)),
        ("Packaging Windows (NSIS + MSI)",                        ("✅ Opérationnel", True, GREEN_OK)),
    ],
    col_widths=[10, 6]
)

h2(doc, "16.2  Résultats des métriques clés")
progress_bar(doc, "Faithfulness (anti-hallucination)",    0.810, "> 0.75")
progress_bar(doc, "Context Relevance (qualité retrieval)", 0.740, "> 0.70")
progress_bar(doc, "Answer Correctness (justesse)",         0.580, "> 0.70",
             "Axe d'amélioration principal — dépend de l'exhaustivité documentaire")
doc.add_paragraph()

h2(doc, "16.3  Apports personnels")
body(doc, "Ce projet m'a permis d'acquérir une expertise concrète en :")
bullet(doc, [
    "Architecture et implémentation complète d'un pipeline RAG en production",
    "Intégration de LLM locaux dans un contexte professionnel contraignant",
    "Développement fullstack asynchrone (FastAPI + React + Tauri + PostgreSQL)",
    "Gestion d'une base documentaire multiformat hétérogène",
    "Évaluation quantitative rigoureuse de systèmes LLM (LLM-as-judge, RAGAS)",
    "Conception UX d'une interface contrainte (environnement hospitalier Windows)",
])

info_box(doc,
    "CHUbot constitue une base solide pour une mise en production étendue. "
    "L'architecture modulaire facilite l'ajout de nouveaux documents, de nouvelles "
    "catégories de questions et de nouvelles fonctionnalités sans refonte majeure.",
    "E9F7EF", "1E8B4C"
)


# ══════════════════════════════════════════════════════════════════════════════
# RÉFÉRENCES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "17. Références et bibliographie")
make_table(doc,
    ["Référence", "Description"],
    [
        ("Lewis et al. (2020)",       '"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks" — papier fondateur du RAG (NeurIPS 2020)'),
        ("Es et al. (2023)",          '"RAGAS: Automated Evaluation of Retrieval Augmented Generation" — framework d\'évaluation standardisé'),
        ("Nussbaum et al. (2024)",    '"Nomic Embed: Training a Reproducible Long Context Text Embedder" — modèle d\'embedding open source'),
        ("LangChain Documentation",   "https://docs.langchain.com — orchestration LLM/RAG"),
        ("ChromaDB Documentation",    "https://docs.trychroma.com — base vectorielle locale"),
        ("Ollama Documentation",      "https://ollama.ai — runtime LLM local"),
        ("FlashRank",                 "Prithiviraj Damodaran (2023) — cross-encoder léger CPU-only"),
        ("Tauri Documentation",       "https://tauri.app — framework desktop Rust/JS"),
        ("FastAPI Documentation",     "https://fastapi.tiangolo.com — framework Python async"),
        ("ms-marco-MiniLM-L-12-v2",   "Microsoft Research — modèle de reranking cross-encoder"),
    ],
    col_widths=[5, 11]
)


# ══════════════════════════════════════════════════════════════════════════════
# ANNEXES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "Annexes")
h2(doc, "Annexe A — Variables d'environnement")
code_block(doc,
"# Application\n"
"APP_NAME=CHUbot\n"
"APP_VERSION=0.1.0\n"
"APP_ENV=production\n"
"DEBUG=false\n\n"
"# LLM (Ollama hébergé CHU)\n"
"OLLAMA_BASE_URL=https://llm.chu-angers.fr/ollama\n"
"OLLAMA_MODEL=llama3.1:latest\n\n"
"# Embeddings (Ollama local)\n"
"EMBEDDING_BASE_URL=http://localhost:11434\n"
"EMBEDDING_MODEL=nomic-embed-text\n\n"
"# Base de données\n"
"DATABASE_URL=postgresql+asyncpg://user:password@localhost:5432/chubot\n\n"
"# RAG\n"
"CHUNK_SIZE=2000\n"
"CHUNK_OVERLAP=400\n"
"RETRIEVAL_TOP_K=5\n"
"VECTOR_STORE_PATH=./data/indexes"
)

h2(doc, "Annexe B — Commandes utiles")
code_block(doc,
"# Démarrer le serveur de développement\n"
"uvicorn bot.api.app:app --host 127.0.0.1 --port 8765 --reload\n\n"
"# Démarrer l'interface Tauri\n"
"cd bot/ui && npx tauri dev\n\n"
"# Indexer tous les documents\n"
"python -m bot.ingestion.indexer --all\n\n"
"# Évaluation standard\n"
"python tests/evaluate.py --limit 10\n\n"
"# Évaluation complète avec RAGAS\n"
"python tests/evaluate.py --ragas\n\n"
"# Vérifier le port 8765\n"
"netstat -ano | findstr :8765"
)

h2(doc, "Annexe C — Structure du projet")
code_block(doc,
"CHUbot/\n"
"├── backend_main.py              # Point d'entrée PyInstaller\n"
"├── requirements.txt             # Dépendances Python\n"
"├── bot/\n"
"│   ├── config/\n"
"│   │   ├── settings.py          # Configuration Pydantic BaseSettings\n"
"│   │   └── departments.py       # 5 départements RH + escalade\n"
"│   ├── ingestion/\n"
"│   │   ├── cleaner.py           # Nettoyage texte\n"
"│   │   ├── loaders.py           # PDF, DOCX, Excel, Web\n"
"│   │   ├── splitter.py          # Chunking + en-têtes contextuels\n"
"│   │   └── indexer.py           # CLI + ChromaDB\n"
"│   ├── retrieval/\n"
"│   │   ├── preprocessor.py      # Acronymes + réécriture\n"
"│   │   ├── query_rewriter.py    # Reformulation LLM\n"
"│   │   ├── reranker.py          # FlashRank\n"
"│   │   ├── prompt.py            # Prompt système RH\n"
"│   │   └── chain.py             # Orchestration RAG complète\n"
"│   ├── api/\n"
"│   │   ├── app.py               # FastAPI + lifespan\n"
"│   │   ├── db/                  # Modèles + schémas + session\n"
"│   │   └── routes/              # chat, documents, feedback, analytics\n"
"│   └── ui/\n"
"│       ├── src/                 # Composants React\n"
"│       └── src-tauri/           # Config Tauri (Rust)\n"
"├── data/\n"
"│   ├── documents/               # ~20 docs RH (PDF, DOCX, Excel)\n"
"│   ├── indexes/                 # ChromaDB persistant\n"
"│   └── eval/                    # 105 questions + résultats CSV\n"
"└── tests/\n"
"    └── evaluate.py              # Framework d'évaluation"
)


# ══════════════════════════════════════════════════════════════════════════════
# SAUVEGARDE
# ══════════════════════════════════════════════════════════════════════════════

output = r"C:\Users\moham\Desktop\CHUbot\rapport_stage.docx"
doc.save(output)
print(f"OK Rapport genere : {output}")
